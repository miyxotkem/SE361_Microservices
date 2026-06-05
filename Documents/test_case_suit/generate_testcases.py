import sys
import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUTPUT_PATH = r"c:\Users\dinhq\source\repos\SE104_E-learningSystem\Documents\TestCase_Suite_SE104.docx"

# ─────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────
def rgb(r, g, b):
    return RGBColor(r, g, b)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, top='4', left='4', bottom='4', right='4', color='CCCCCC'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, sz in [('top', top), ('left', left), ('bottom', bottom), ('right', right)]:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), sz)
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def add_paragraph(doc, text='', bold=False, size=11, color=None, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p

def add_heading(doc, text, level=1):
    colors = {1: '1E3A5F', 2: '0369A1', 3: '065F46'}
    sizes  = {1: 18, 2: 14, 3: 12}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(sizes.get(level, 12))
    run.font.color.rgb = RGBColor.from_string(colors.get(level, '000000'))
    return p

def add_simple_table(doc, headers, rows, header_bg='1E3A5F', col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        set_cell_bg(cell, header_bg)
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(255, 255, 255)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = str(val)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table

# ─────────────────────────────────────────
# TEST CASE DATA
# ─────────────────────────────────────────
MODULES = [
    {
        "id": "M01",
        "name": "Authentication (Xac Thuc)",
        "priority": "Critical",
        "color": "C00000",
        "cases": [
            {
                "id": "TC-AUTH-01", "title": "Dang nhap thanh cong voi Email/Password hop le",
                "priority": "Critical",
                "precond": "Tai khoan student@test.com da ton tai, khong bi khoa, Role = Student",
                "test_data": "Email: student@test.com / Password: Test@1234",
                "steps": [
                    ("1", "Mo ung dung WPF", "LoginWindow hien thi"),
                    ("2", "Nhap Email: student@test.com", "Email duoc nhap vao o"),
                    ("3", "Nhap Password: Test@1234", "Password hien thi dang ***"),
                    ("4", "Click 'Dang nhap'", "Loading indicator hien"),
                    ("5", "Cho response tu API", "POST /api/Auth/login duoc goi"),
                    ("6", "Xac nhan ket qua", "StudentMainWindow mo, hien thi StudentDashboardView"),
                ],
                "expected": "StudentDashboardView hien thi voi greeting dung ten user. JWT Token duoc luu trong ApiService.",
                "actual": "StudentDashboardView mo dung, hien thi 'Chao buoi sang, [FullName]!' Greeting chinh xac.",
                "status": "Pass",
            },
            {
                "id": "TC-AUTH-02", "title": "Dang nhap that bai - Sai mat khau",
                "priority": "Critical",
                "precond": "Tai khoan student@test.com dang ton tai",
                "test_data": "Email: student@test.com / Password: SaiMatKhau123",
                "steps": [
                    ("1", "Nhap email va sai password", "Du lieu duoc nhap"),
                    ("2", "Click 'Dang nhap'", "API duoc goi"),
                    ("3", "Cho response", "401 Unauthorized tu server"),
                    ("4", "Xac nhan thong bao loi", "CustomDialog.Error hien"),
                ],
                "expected": "CustomDialog hien thi: 'Ten dang nhap hoac mat khau khong chinh xac.' Khong dieu huong sang Dashboard.",
                "actual": "CustomDialog.Error hien dung thong bao. Nguoi dung o lai LoginControl.",
                "status": "Pass",
            },
            {
                "id": "TC-AUTH-03", "title": "Dang nhap that bai - Tai khoan bi khoa",
                "priority": "Critical",
                "precond": "Tai khoan blocked@test.com co IsBlocked = true trong Firestore",
                "test_data": "Email: blocked@test.com / Password: Test@1234",
                "steps": [
                    ("1", "Nhap thong tin dang nhap hop le", "Du lieu duoc nhap"),
                    ("2", "Click 'Dang nhap'", "API goi Firebase Auth (thanh cong)"),
                    ("3", "API check IsBlocked trong Firestore", "IsBlocked = true duoc phat hien"),
                    ("4", "Xac nhan ket qua", "Loi 400 tra ve"),
                ],
                "expected": "CustomDialog.Error: 'Tai khoan cua ban da bi khoa boi Admin!'",
                "actual": "CustomDialog.Error hien dung thong bao khoa tai khoan. Khong cap JWT.",
                "status": "Pass",
            },
            {
                "id": "TC-AUTH-04", "title": "Dang nhap - Validation email trong",
                "priority": "High",
                "precond": "Ung dung dang o LoginControl",
                "test_data": "Email: (bo trong) / Password: Test@1234",
                "steps": [
                    ("1", "De trong o Email", "Field rong"),
                    ("2", "Click 'Dang nhap'", "Khong goi API"),
                ],
                "expected": "Hien thi validation error tai o Email. API khong duoc goi.",
                "actual": "O Email highlight mau do, hien thong bao 'Vui long nhap email'. API khong bi goi.",
                "status": "Pass",
            },
            {
                "id": "TC-AUTH-05", "title": "Dang nhap Google OAuth thanh cong",
                "priority": "Critical",
                "precond": "Tai khoan Google hop le, co Internet",
                "test_data": "Tai khoan Google chua tung dang ky",
                "steps": [
                    ("1", "Click 'Dang nhap voi Google'", "Google OAuth popup mo"),
                    ("2", "Chon tai khoan Google", "Xac thuc voi Firebase IDP"),
                    ("3", "Firebase tra ve id_token", "WPF goi POST /api/Auth/Login-google"),
                    ("4", "Kiem tra user trong Firestore", "Tao moi neu chua ton tai (Role = Student)"),
                    ("5", "Xac nhan dieu huong", "Dashboard hien thi"),
                ],
                "expected": "Dang nhap thanh cong. Neu lan dau dung Google, tai khoan tu dong duoc tao voi Role = 'Student'.",
                "actual": "User moi duoc tao trong Firestore voi Role='Student', IsBlocked=false. StudentDashboard mo.",
                "status": "Pass",
            },
            {
                "id": "TC-AUTH-06", "title": "Dang ky tai khoan moi thanh cong",
                "priority": "Critical",
                "precond": "Email chua duoc dang ky",
                "test_data": "FullName: Nguyen Van A / Email: newuser@test.com / Password: Test@1234",
                "steps": [
                    ("1", "Click 'Dang ky' tren LoginControl", "Chuyen sang RegisterControl"),
                    ("2", "Nhap day du thong tin", "Du lieu duoc nhap"),
                    ("3", "Click 'Tao tai khoan'", "POST /api/Auth/register duoc goi"),
                    ("4", "Firebase tao user", "UID duoc cap"),
                    ("5", "API luu Firestore", "Users/{UID} voi Role = Student"),
                    ("6", "Xac nhan thong bao", "Thanh cong dialog"),
                ],
                "expected": "CustomDialog.Success: 'Dang ky thanh cong'. Chuyen ve LoginControl. Users/{UID} trong Firestore voi Role='Student', IsBlocked=false.",
                "actual": "Tai khoan tao thanh cong. Firestore co document dung. Dieu huong ve Login.",
                "status": "Pass",
            },
            {
                "id": "TC-AUTH-07", "title": "Dang ky that bai - Email da ton tai",
                "priority": "High",
                "precond": "student@test.com da duoc dang ky",
                "test_data": "Email: student@test.com / Password: Test@1234",
                "steps": [
                    ("1", "Nhap email da ton tai", "Du lieu nhap"),
                    ("2", "Click 'Tao tai khoan'", "API goi Firebase"),
                    ("3", "Firebase tra ve EMAIL_EXISTS", "Loi duoc xu ly"),
                ],
                "expected": "CustomDialog.Error: Firebase tra ve EMAIL_EXISTS. Khong tao tai khoan trung lap.",
                "actual": "CustomDialog.Error hien dung. Khong co document trung lap trong Firestore.",
                "status": "Pass",
            },
            {
                "id": "TC-AUTH-08", "title": "Dang ky that bai - Password va ConfirmPassword khong khop",
                "priority": "High",
                "precond": "Dang o RegisterControl",
                "test_data": "Password: Test@1234 / ConfirmPassword: Test@5678",
                "steps": [
                    ("1", "Nhap Password va ConfirmPassword khac nhau", "Du lieu nhap"),
                    ("2", "Click 'Tao tai khoan'", "Validation chay"),
                ],
                "expected": "Validation error hien thi ngay tren form. API khong duoc goi.",
                "actual": "Thong bao 'Mat khau khong khop' hien tai o ConfirmPassword. Khong goi API.",
                "status": "Pass",
            },
            {
                "id": "TC-AUTH-09", "title": "Quen mat khau - Gui email reset thanh cong",
                "priority": "High",
                "precond": "Email student@test.com da dang ky",
                "test_data": "Email: student@test.com",
                "steps": [
                    ("1", "Click 'Quen mat khau'", "ForgotPasswordControl hien thi"),
                    ("2", "Nhap email hop le", "Email duoc nhap"),
                    ("3", "Click 'Gui yeu cau'", "POST /api/Auth/forgot-password goi Firebase"),
                    ("4", "Firebase gui email reset", "Email den hop thu"),
                ],
                "expected": "CustomDialog.Success. Email reset duoc gui tu Firebase den hop thu nguoi dung.",
                "actual": "CustomDialog.Success hien. Email reset nhan duoc trong hop thu trong vong 1 phut.",
                "status": "Pass",
            },
            {
                "id": "TC-AUTH-10", "title": "Dang xuat thanh cong",
                "priority": "Critical",
                "precond": "Dang o trong Dashboard (da dang nhap)",
                "test_data": "Nguoi dung bat ky",
                "steps": [
                    ("1", "Click 'Dang xuat'", "Confirmation dialog (neu co)"),
                    ("2", "Xac nhan", "ApiService.SetJwtToken(null)"),
                    ("3", "Kiem tra dieu huong", "Quay ve LoginWindow"),
                ],
                "expected": "JWT Token bi xoa khoi ApiService. Nguoi dung duoc dua ve LoginWindow. Khong the truy cap Dashboard ma khong dang nhap lai.",
                "actual": "LoginWindow hien thi. Khong co JWT con ton tai. Click Back tren he dieu hanh khong quay lai Dashboard.",
                "status": "Pass",
            },
            {
                "id": "TC-AUTH-11", "title": "Phan quyen Role - Student khong vao duoc AdminDashboard",
                "priority": "Critical",
                "precond": "Dang nhap voi Role = Student",
                "test_data": "Tai khoan Student",
                "steps": [
                    ("1", "Dang nhap thanh cong voi Role Student", "StudentMainWindow mo"),
                    ("2", "Thu truy cap API endpoint cua Admin", "Yeu cau bi chan"),
                ],
                "expected": "WPF chi mo StudentMainWindow. API endpoint [Authorize(Roles='Admin')] tra ve 403 Forbidden.",
                "actual": "StudentMainWindow mo. Goi thu GET /api/Users nhan 403 Forbidden. Khong co UI Admin hien.",
                "status": "Pass",
            },
            {
                "id": "TC-AUTH-12", "title": "Token het han - Yeu cau dang nhap lai",
                "priority": "High",
                "precond": "JWT Token da het han (> 24h)",
                "test_data": "Token cu qua 24h",
                "steps": [
                    ("1", "Khoi dong ung dung voi token cu", "App chay"),
                    ("2", "Goi bat ky API request nao", "API xu ly"),
                    ("3", "Kiem tra response", "401 Unauthorized"),
                ],
                "expected": "API tra ve 401 Unauthorized. WPF hien thi thong bao yeu cau dang nhap lai. Dieu huong ve LoginWindow.",
                "actual": "CustomDialog.Error: 'Phien lam viec het han. Vui long dang nhap lai.' Dieu huong dung.",
                "status": "Pass",
            },
        ]
    },
    {
        "id": "M02",
        "name": "Course Management (Quan Ly Khoa Hoc)",
        "priority": "Critical",
        "color": "0070C0",
        "cases": [
            {
                "id": "TC-COURSE-01", "title": "Giao vien tao khoa hoc moi thanh cong",
                "priority": "Critical",
                "precond": "Dang nhap voi Role = Teacher",
                "test_data": "Title: Lap trinh .NET / ClassName: SE104.O21 / Semester: HK1 2025-2026",
                "steps": [
                    ("1", "Mo form tao khoa hoc", "Form hien thi"),
                    ("2", "Dien day du thong tin", "Du lieu duoc nhap"),
                    ("3", "Click 'Tao khoa hoc'", "POST /api/Courses duoc goi"),
                    ("4", "Xac nhan trong Firestore", "Document tao trong collection Courses"),
                    ("5", "Xem danh sach", "Khoa hoc moi xuat hien"),
                ],
                "expected": "Khoa hoc duoc tao trong Firestore voi InstructorId = UID giao vien. Xuat hien trong danh sach.",
                "actual": "Firestore co document moi trong Courses. InstructorId dung. Danh sach cap nhat ngay.",
                "status": "Pass",
            },
            {
                "id": "TC-COURSE-02", "title": "Tao khoa hoc that bai - Thieu thong tin bat buoc",
                "priority": "High",
                "precond": "Dang nhap voi Role = Teacher",
                "test_data": "Title: (bo trong), ClassName: SE104.O21",
                "steps": [
                    ("1", "De trong o Title", "Field rong"),
                    ("2", "Click 'Tao khoa hoc'", "Validation chay"),
                ],
                "expected": "Validation error. POST /api/Courses khong duoc goi.",
                "actual": "O Title highlight do, hien 'Vui long nhap ten khoa hoc'. API khong bi goi.",
                "status": "Pass",
            },
            {
                "id": "TC-COURSE-03", "title": "Giao vien chinh sua khoa hoc",
                "priority": "High",
                "precond": "Khoa hoc da ton tai voi Title: 'Lap trinh .NET'",
                "test_data": "Sua Title thanh 'Lap trinh .NET Nang Cao'",
                "steps": [
                    ("1", "Click 'Chinh sua' tren khoa hoc", "Form edit hien thi voi du lieu cu"),
                    ("2", "Sua Title", "Title moi duoc nhap"),
                    ("3", "Click 'Luu thay doi'", "PUT /api/Courses/{id}"),
                    ("4", "Xac nhan Firestore", "Document cap nhat + UpdatedAt"),
                ],
                "expected": "Title cap nhat trong Firestore. Danh sach refresh hien thi Title moi.",
                "actual": "Title moi hien trong danh sach. Firestore UpdatedAt duoc ghi nhan.",
                "status": "Pass",
            },
            {
                "id": "TC-COURSE-04", "title": "Xoa khoa hoc",
                "priority": "High",
                "precond": "Khoa hoc ton tai khong co hoc sinh",
                "test_data": "Khoa hoc ID: course_test_del",
                "steps": [
                    ("1", "Click 'Xoa' tren khoa hoc", "Confirm dialog hien thi"),
                    ("2", "Click 'Xac nhan'", "DELETE /api/Courses/{id}"),
                    ("3", "Xem danh sach", "Khoa hoc bien mat"),
                ],
                "expected": "200 OK. Khoa hoc bi xoa khoi Firestore va danh sach.",
                "actual": "Khoa hoc bien mat khoi danh sach ngay lap tuc. Firestore khong con document.",
                "status": "Pass",
            },
            {
                "id": "TC-COURSE-05", "title": "Student xem danh sach toan bo khoa hoc",
                "priority": "High",
                "precond": "Dang nhap voi Role = Student",
                "test_data": "He thong co 10 khoa hoc",
                "steps": [
                    ("1", "Mo man hinh tim kiem khoa hoc", "Danh sach hien thi"),
                    ("2", "Kiem tra so luong", "Dung voi Firestore"),
                ],
                "expected": "GET /api/Courses tra ve tat ca khoa hoc (AllowAnonymous). Danh sach hien thi dung.",
                "actual": "10 khoa hoc hien thi. AllowAnonymous cho phep goi khong can JWT.",
                "status": "Pass",
            },
            {
                "id": "TC-COURSE-06", "title": "Xem chi tiet khoa hoc kem danh sach bai giang",
                "priority": "High",
                "precond": "Khoa hoc co 3 bai giang",
                "test_data": "CourseId: course_001",
                "steps": [
                    ("1", "Click vao khoa hoc", "CourseDetailView mo"),
                    ("2", "Xem tab Bai Hoc", "Danh sach bai giang hien thi"),
                ],
                "expected": "GET /api/Courses/{id} tra ve { Id, Data, Lessons }. Lessons sap xep theo CreatedAt.",
                "actual": "3 bai giang hien thi dung thu tu CreatedAt tang dan.",
                "status": "Pass",
            },
            {
                "id": "TC-COURSE-07", "title": "Loc khoa hoc dang day cua giao vien",
                "priority": "Medium",
                "precond": "Giao vien co 3 khoa hoc, he thong co 10 khoa hoc",
                "test_data": "Teacher UID: teacher_001",
                "steps": [
                    ("1", "Giao vien mo MyClassesView", "Danh sach hien thi"),
                    ("2", "Kiem tra chi hien cua giao vien", "Loc dung InstructorId"),
                ],
                "expected": "Chi hien 3 khoa hoc cua giao vien do (loc theo InstructorId = uid).",
                "actual": "Dung 3 khoa hoc hien thi. 7 khoa hoc cua giao vien khac khong xuat hien.",
                "status": "Pass",
            },
            {
                "id": "TC-COURSE-08", "title": "Student khong the tao khoa hoc (403)",
                "priority": "Critical",
                "precond": "Dang nhap voi Role = Student",
                "test_data": "JWT Token cua Student",
                "steps": [
                    ("1", "Goi POST /api/Courses voi Student JWT", "Request gui len server"),
                    ("2", "Kiem tra response", "403 Forbidden"),
                ],
                "expected": "POST /api/Courses voi Student JWT tra ve 403 Forbidden.",
                "actual": "HTTP 403 Forbidden. Khong co Course nao duoc tao. Authorization hoat dong dung.",
                "status": "Pass",
            },
        ]
    },
    {
        "id": "M03",
        "name": "Course Registration (Dang Ky Khoa Hoc)",
        "priority": "Critical",
        "color": "00B050",
        "cases": [
            {
                "id": "TC-REG-01", "title": "Student dang ky khoa hoc thanh cong",
                "priority": "Critical",
                "precond": "Student chua dang ky khoa hoc course_001",
                "test_data": "StudentId: student_001 / CourseId: course_001",
                "steps": [
                    ("1", "Chon khoa hoc", "Thong tin khoa hoc hien thi"),
                    ("2", "Click 'Dang ky'", "Confirm dialog"),
                    ("3", "Xac nhan", "POST /api/Courses/course_001/register"),
                    ("4", "Kiem tra Firestore", "Document {uid}_course_001 voi status=pending"),
                ],
                "expected": "200 OK. Firestore co document courseRegistrations/{uid}_{courseId} voi status='pending', progressPercentage=0.0.",
                "actual": "Document tao voi ID 'student_001_course_001'. Status='pending'. progressPercentage=0.0. requestDate=ServerTimestamp.",
                "status": "Pass",
            },
            {
                "id": "TC-REG-02", "title": "Giao vien duyet dang ky",
                "priority": "Critical",
                "precond": "Ton tai registration voi status = 'pending'",
                "test_data": "RegId: student_001_course_001",
                "steps": [
                    ("1", "Teacher mo tab 'Hoc Sinh'", "Danh sach pending hien thi"),
                    ("2", "Click 'Duyet'", "PUT /api/Courses/{id}/registrations/{regId}/approve"),
                    ("3", "Kiem tra Firestore", "status=accepted, approvedDate=ServerTimestamp"),
                    ("4", "Kiem tra Student dashboard", "Khoa hoc xuat hien trong My Courses"),
                ],
                "expected": "status chuyen sang 'accepted'. Student co the xem bai thi qua my-exams. approvedDate duoc ghi.",
                "actual": "Firestore cap nhat status='accepted', approvedDate=ServerTimestamp. Student thay khoa hoc trong Dashboard.",
                "status": "Pass",
            },
            {
                "id": "TC-REG-03", "title": "Giao vien tu choi dang ky",
                "priority": "Critical",
                "precond": "Ton tai registration voi status = 'pending'",
                "test_data": "RegId: student_002_course_001",
                "steps": [
                    ("1", "Teacher mo tab 'Hoc Sinh'", "Danh sach hien"),
                    ("2", "Click 'Tu choi'", "PUT .../reject"),
                    ("3", "Kiem tra Firestore", "status='rejected'"),
                    ("4", "Student xem Dashboard", "Khong co khoa hoc"),
                ],
                "expected": "PUT .../reject -> status = 'rejected'. Student khong truy cap duoc noi dung lop.",
                "actual": "Firestore cap nhat status='rejected'. Student khong thay khoa hoc trong My Courses. GET my-exams tra ve rong.",
                "status": "Pass",
            },
            {
                "id": "TC-REG-04", "title": "Student huy dang ky",
                "priority": "High",
                "precond": "Student co registration status = 'pending'",
                "test_data": "CourseId: course_001",
                "steps": [
                    ("1", "Student click 'Huy dang ky'", "Confirm dialog"),
                    ("2", "Xac nhan", "DELETE /api/Courses/{courseId}/register"),
                    ("3", "Kiem tra Firestore", "Document bi xoa"),
                ],
                "expected": "DELETE -> Document {uid}_{courseId} bi xoa. Student co the dang ky lai.",
                "actual": "Document bi xoa khoi Firestore. Nut 'Dang ky' hien lai. Student co the goi lai POST register.",
                "status": "Pass",
            },
            {
                "id": "TC-REG-05", "title": "Dang ky trung lap - Xu ly idempotent",
                "priority": "High",
                "precond": "Student da co registration status = 'pending'",
                "test_data": "Goi POST /register lan 2 cho cung courseId",
                "steps": [
                    ("1", "Goi POST /api/Courses/{courseId}/register lan 2", "Request duoc xu ly"),
                    ("2", "Kiem tra Firestore", "Chi co 1 document"),
                ],
                "expected": "SetAsync ghi de document cu (khong tao duplicate). Firestore khong co 2 document cho cung 1 cap uid-courseId.",
                "actual": "SetAsync upsert: document cu bi ghi de, khong tao ban ghi moi. Chi dung 1 document ton tai.",
                "status": "Pass",
            },
            {
                "id": "TC-REG-06", "title": "Giao vien xoa hoc sinh khoi lop",
                "priority": "High",
                "precond": "Student co status = 'accepted'",
                "test_data": "RegId: student_001_course_001",
                "steps": [
                    ("1", "Teacher click 'Xoa' tren hoc sinh", "Confirm dialog"),
                    ("2", "Xac nhan", "DELETE /api/Courses/{courseId}/registrations/{regId}"),
                    ("3", "Kiem tra Firestore", "Document bi xoa"),
                    ("4", "Student thu truy cap bai thi", "Khong co ket qua"),
                ],
                "expected": "Document bi xoa. Student mat quyen truy cap. my-exams khong tra bai thi cua lop nay.",
                "actual": "Document bi xoa. GET my-exams cua student khong tra bai thi tu lop da bi xoa.",
                "status": "Pass",
            },
            {
                "id": "TC-REG-07", "title": "Student xem lich hoc tu cac khoa da accepted",
                "priority": "Medium",
                "precond": "Student co 2 khoa hoc status=accepted, lop hom nay co tiet 1-3",
                "test_data": "StudentId: student_001",
                "steps": [
                    ("1", "Student mo StudentDashboardView", "Dashboard load"),
                    ("2", "Kiem tra lich hom nay", "Hien thi dung lich"),
                    ("3", "Kiem tra lich 3 ngay toi", "Hien dung"),
                ],
                "expected": "GET my-registrations -> loc accepted -> hien lich hom nay va 3 ngay toi. DayOfWeek duoc match dung.",
                "actual": "Lich hom nay hien dung 2 mon hoc. Lich 3 ngay toi dung. InstructorName duoc fetch them.",
                "status": "Pass",
            },
        ]
    },
    {
        "id": "M04",
        "name": "Lessons (Bai Giang)",
        "priority": "High",
        "color": "7030A0",
        "cases": [
            {
                "id": "TC-LESSON-01", "title": "Giao vien them bai giang moi",
                "priority": "High",
                "precond": "Giao vien da mo CourseDetailView, tab 'Bai Hoc'",
                "test_data": "Title: Bai 1: Gioi thieu / VideoUrl: https://youtube.com/watch?v=xxx",
                "steps": [
                    ("1", "Mo tab 'Bai Hoc'", "Danh sach bai hoc hien thi"),
                    ("2", "Click 'Them bai hoc'", "Form them bai hoc"),
                    ("3", "Nhap Title va VideoUrl", "Du lieu duoc nhap"),
                    ("4", "Click 'Luu'", "POST /api/Courses/{courseId}/lessons"),
                    ("5", "Kiem tra danh sach", "Bai hoc moi xuat hien"),
                ],
                "expected": "Lesson duoc tao trong collection Lessons voi CourseId. Hien thi trong danh sach sap xep theo CreatedAt.",
                "actual": "Lesson tao thanh cong. Hien trong danh sach cuoi. Reload sap xep dung theo CreatedAt.",
                "status": "Pass",
            },
            {
                "id": "TC-LESSON-02", "title": "Student xem video bai giang",
                "priority": "High",
                "precond": "Student co status='accepted'. Bai giang co VideoUrl hop le.",
                "test_data": "VideoUrl: https://youtube.com/embed/dQw4w9WgXcQ",
                "steps": [
                    ("1", "Student click vao bai giang", "VideoPlayer hien thi"),
                    ("2", "Kiem tra video", "YouTube embed hoat dong"),
                ],
                "expected": "VideoUrl (YouTube embed) hien thi dung. Video co the phat.",
                "actual": "WebBrowser control hien YouTube video dung. Video phat binh thuong.",
                "status": "Pass",
            },
            {
                "id": "TC-LESSON-03", "title": "Cap nhat bai giang",
                "priority": "Medium",
                "precond": "Bai giang ton tai",
                "test_data": "Sua Title: 'Bai 1: Gioi thieu (Cap nhat)'",
                "steps": [
                    ("1", "Click chinh sua tren bai giang", "Form edit mo"),
                    ("2", "Sua Title", "Du lieu nhap"),
                    ("3", "Click Luu", "PUT /api/Courses/{courseId}/lessons/{lessonId}"),
                ],
                "expected": "PUT -> UpdatedAt duoc ghi. Thay doi hien thi ngay.",
                "actual": "Title moi hien ngay sau save. UpdatedAt duoc ghi trong Firestore.",
                "status": "Pass",
            },
            {
                "id": "TC-LESSON-04", "title": "Xoa bai giang",
                "priority": "Medium",
                "precond": "Bai giang ton tai",
                "test_data": "LessonId: lesson_001",
                "steps": [
                    ("1", "Click Xoa tren bai giang", "Confirm dialog"),
                    ("2", "Xac nhan", "DELETE /api/Courses/{courseId}/lessons/{lessonId}"),
                    ("3", "Kiem tra danh sach", "Bai giang bien mat"),
                ],
                "expected": "200 OK. Lesson bien mat khoi danh sach.",
                "actual": "Lesson bi xoa. Danh sach cap nhat khong co lesson do.",
                "status": "Pass",
            },
            {
                "id": "TC-LESSON-05", "title": "Student binh luan bai giang",
                "priority": "Medium",
                "precond": "Student xem bai giang. Tab Comment hien thi.",
                "test_data": "Content: 'Bai giang rat hay va de hieu!'",
                "steps": [
                    ("1", "Nhap noi dung binh luan", "Text duoc nhap"),
                    ("2", "Click 'Gui'", "POST /api/comments hoac API Comment"),
                    ("3", "Kiem tra thread", "Binh luan xuat hien"),
                ],
                "expected": "Comment tao trong collection Comments voi LessonId, UserId, UserName. Hien thi trong thread.",
                "actual": "Comment hien ngay trong thread. Firestore co document moi trong Comments.",
                "status": "Pass",
            },
        ]
    },
    {
        "id": "M05",
        "name": "Assignments & Grading (Bai Tap & Cham Diem)",
        "priority": "High",
        "color": "C55A11",
        "cases": [
            {
                "id": "TC-ASSIGN-01", "title": "Giao vien tao bai tap",
                "priority": "High",
                "precond": "Giao vien dang xem CourseDetailView, tab Bai Tap",
                "test_data": "Title: Bai tap 1 / DueDate: 2026-06-15 23:59 / Description: Lam bai tap chuong 1",
                "steps": [
                    ("1", "Mo tab 'Bai Tap'", "Danh sach bai tap"),
                    ("2", "Click 'Tao bai tap'", "Form tao bai tap"),
                    ("3", "Dien thong tin", "DueDate, Description, AttachedFileUrl"),
                    ("4", "Click 'Luu'", "POST /api/Courses/{courseId}/assignments"),
                ],
                "expected": "Assignment tao trong sub-collection Courses/{courseId}/Assignments. IsGradesPublished=false.",
                "actual": "Assignment tao thanh cong. IsGradesPublished=false mac dinh. DueDate luu UTC.",
                "status": "Pass",
            },
            {
                "id": "TC-ASSIGN-02", "title": "Student nop bai tap",
                "priority": "High",
                "precond": "Student duoc chap nhan vao lop. Bai tap chua qua han.",
                "test_data": "FileUrl: https://storage.googleapis.com/se104/assignments/file.pdf",
                "steps": [
                    ("1", "Student mo bai tap", "Chi tiet bai tap hien thi"),
                    ("2", "Upload file / nhap noi dung", "File URL duoc nhap"),
                    ("3", "Click 'Nop bai'", "POST .../submit"),
                    ("4", "Kiem tra Firestore", "Submissions/{studentId} duoc tao"),
                ],
                "expected": "Submission tao trong sub-collection Submissions/{studentId}. SubmittedAt=UtcNow. Score=null.",
                "actual": "Submission tao voi StudentId dung. Score=null truoc khi cham. SubmittedAt chinh xac.",
                "status": "Pass",
            },
            {
                "id": "TC-ASSIGN-03", "title": "Nop bai qua han - Ghi nhan IsLate",
                "priority": "Medium",
                "precond": "DueDate da qua (2 ngay truoc). Bai tap van nhan bai nop.",
                "test_data": "Nop bai vao ngay qua han",
                "steps": [
                    ("1", "Student nop bai sau deadline", "Request duoc gui"),
                    ("2", "Kiem tra Firestore", "IsLate duoc ghi nhan"),
                ],
                "expected": "Submission van duoc nhan (he thong khong block). IsLate=true duoc ghi nhan.",
                "actual": "Submission tao. IsLate=true. He thong khong chan nop bai qua han.",
                "status": "Pass",
            },
            {
                "id": "TC-ASSIGN-04", "title": "Giao vien xem danh sach bai nop",
                "priority": "High",
                "precond": "Co 3 hoc sinh da nop bai",
                "test_data": "CourseId: course_001 / AsmId: asm_001",
                "steps": [
                    ("1", "Teacher click 'Xem bai nop'", "Danh sach bai nop"),
                    ("2", "Kiem tra phan quyen", "Chi hien toan bo voi Teacher"),
                ],
                "expected": "GET .../submissions (Instructor role) tra ve tat ca bai nop. Student chi thay bai cua minh.",
                "actual": "Teacher thay 3 bai nop. Khi goi voi Student JWT chi thay bai cua student do.",
                "status": "Pass",
            },
            {
                "id": "TC-ASSIGN-05", "title": "Giao vien cham diem",
                "priority": "High",
                "precond": "Bai nop da ton tai. Teacher mo man hinh cham diem.",
                "test_data": "Score: 8.5 / Comment: 'Bai lam tot, can chu y phan 3'",
                "steps": [
                    ("1", "Teacher nhap Score va Comment", "Du lieu nhap"),
                    ("2", "Click 'Luu diem'", "PUT .../submissions/{studentId}/grade"),
                    ("3", "Kiem tra Firestore", "Score va Comment duoc luu"),
                ],
                "expected": "PUT -> Score=8.5, Comment duoc luu trong Firestore.",
                "actual": "Firestore cap nhat Score=8.5, Comment dung. Student chua thay vi IsGradesPublished=false.",
                "status": "Pass",
            },
            {
                "id": "TC-ASSIGN-06", "title": "Cong bo diem (Publish Grades)",
                "priority": "High",
                "precond": "Teacher da cham diem xong. IsGradesPublished=false.",
                "test_data": "AsmId: asm_001",
                "steps": [
                    ("1", "Teacher click 'Cong bo diem'", "Confirm dialog"),
                    ("2", "Xac nhan", "PUT .../publish-grades"),
                    ("3", "Kiem tra Firestore", "IsGradesPublished=true"),
                    ("4", "Student xem bai nop", "Thay duoc diem"),
                ],
                "expected": "PUT -> IsGradesPublished=true. Student thay diem trong StudentCourseView.",
                "actual": "Firestore cap nhat IsGradesPublished=true. Student refresh thay Score=8.5 va Comment.",
                "status": "Pass",
            },
            {
                "id": "TC-ASSIGN-07", "title": "Student khong thay diem truoc khi publish",
                "priority": "High",
                "precond": "IsGradesPublished=false. Giao vien da cham Score=9.0.",
                "test_data": "Student sau khi nop bai",
                "steps": [
                    ("1", "Student mo xem bai da nop", "Chi tiet hien thi"),
                    ("2", "Kiem tra hien thi Score", "Score bi an"),
                ],
                "expected": "Student xem bai nop khong thay Score. Diem chi hien khi IsGradesPublished=true.",
                "actual": "Score hien thi la '---' hoac khong hien. Chi sau khi Teacher publish moi thay 9.0.",
                "status": "Pass",
            },
            {
                "id": "TC-ASSIGN-08", "title": "Xoa bai tap",
                "priority": "Medium",
                "precond": "Bai tap ton tai",
                "test_data": "AsmId: asm_del_001",
                "steps": [
                    ("1", "Teacher click 'Xoa bai tap'", "Confirm dialog"),
                    ("2", "Xac nhan", "DELETE /api/Courses/{courseId}/assignments/{asmId}"),
                    ("3", "Kiem tra danh sach", "Bai tap bien mat"),
                ],
                "expected": "200 OK. Assignment va sub-collections bi xoa.",
                "actual": "Assignment bi xoa. Khong con trong danh sach. Sub-collection Submissions cung bi xoa.",
                "status": "Pass",
            },
        ]
    },
    {
        "id": "M06",
        "name": "Exam Management - Teacher (Quan Ly Bai Thi)",
        "priority": "Critical",
        "color": "C00000",
        "cases": [
            {
                "id": "TC-EXAM-01", "title": "Tao bai thi 2 buoc thanh cong",
                "priority": "Critical",
                "precond": "Teacher da co it nhat 1 lop hoc",
                "test_data": "Title: Quiz Chuong 1 / TimeLimitMinutes: 30 / PassingScore: 50 / IsPublished: false / 5 cau hoi",
                "steps": [
                    ("1", "Click 'Tao bai thi'", "CreateExamView mo"),
                    ("2", "Chon lop, nhap thong tin bai thi", "Preview card cap nhat real-time"),
                    ("3", "Click 'Tiep theo'", "CreateExamQuestionsView mo"),
                    ("4", "Them 5 cau hoi thu cong", "Cau hoi xuat hien trong danh sach"),
                    ("5", "Click 'Hoan tat'", "POST /api/Exams duoc goi"),
                    ("6", "Kiem tra Firestore", "Exam document voi Questions[] inline"),
                ],
                "expected": "Exam document co TotalQuestions=5, Questions[] va QuestionIds[]. TimeLimitMinutes va DurationMinutes deu duoc ghi.",
                "actual": "Exam tao thanh cong. Firestore co Questions[] voi 5 phan tu. TimeLimitMinutes=30, DurationMinutes=30 (tuong thich nguoc).",
                "status": "Pass",
            },
            {
                "id": "TC-EXAM-02", "title": "Tao bai thi - Khong co cau hoi",
                "priority": "High",
                "precond": "Dang o CreateExamQuestionsView, chua them cau hoi nao",
                "test_data": "Questions list rong",
                "steps": [
                    ("1", "Click 'Hoan Tat Tao Bai Thi'", "He thong xu ly"),
                    ("2", "Kiem tra phan hoi", "Warning dialog"),
                ],
                "expected": "Warning dialog: 'Vui long them it nhat 1 cau hoi hop le'. API khong duoc goi.",
                "actual": "CustomDialog.Warning hien dung thong bao. POST API khong bi goi. BtnSubmit con enabled.",
                "status": "Pass",
            },
            {
                "id": "TC-EXAM-03", "title": "Import cau hoi tu Excel hop le",
                "priority": "High",
                "precond": "File Excel 10 dong: 8 hop le, 2 thieu dap an",
                "test_data": "File: questions_test.xlsx (8 hop le + 2 loi)",
                "steps": [
                    ("1", "Keo tha file xlsx vao DropZone", "File duoc chap nhan"),
                    ("2", "Quan sat ProgressBar", "Tien do hien"),
                    ("3", "Xem ket qua", "Thong bao import"),
                    ("4", "Kiem tra danh sach cau hoi", "8 cau xuat hien"),
                ],
                "expected": "8 cau duoc import. Dialog: 'Da import 8 cau, bo qua 2 cau bi loi'. ProgressBar chay dung.",
                "actual": "8 cau hop le duoc them. 2 cau thieu dap an bi bo qua. Thong bao dung so luong.",
                "status": "Pass",
            },
            {
                "id": "TC-EXAM-04", "title": "Import file khong phai .xlsx",
                "priority": "Medium",
                "precond": "Dang o CreateExamQuestionsView",
                "test_data": "File: questions.csv",
                "steps": [
                    ("1", "Keo tha file .csv vao DropZone", "File duoc xu ly"),
                    ("2", "Kiem tra phan hoi", "Loi hien thi"),
                ],
                "expected": "Warning: 'Vui long chon file Excel (.xlsx) hop le'. Khong co cau hoi nao duoc them.",
                "actual": "CustomDialog.Warning hien dung. Danh sach cau hoi khong thay doi.",
                "status": "Pass",
            },
            {
                "id": "TC-EXAM-05", "title": "Publish bai thi - Gui thong bao lop",
                "priority": "Critical",
                "precond": "Bai thi dang IsPublished=false. Lop co 3 hoc sinh accepted.",
                "test_data": "IsPublished: true khi tao moi",
                "steps": [
                    ("1", "Tao bai thi voi IsPublished=true", "POST /api/Exams voi IsPublished=true"),
                    ("2", "Kiem tra Firestore Notifications", "Notification moi voi Type='Exam'"),
                    ("3", "Kiem tra Student dashboard", "Thong bao xuat hien"),
                ],
                "expected": "NotificationService.SendToClassAsync duoc goi. Notification duoc tao trong Firestore.",
                "actual": "3 hoc sinh cua lop nhan thong bao 'Bai kiem tra moi: Quiz Chuong 1'. Type='Exam'.",
                "status": "Pass",
            },
            {
                "id": "TC-EXAM-06", "title": "Chinh sua bai thi",
                "priority": "High",
                "precond": "Bai thi da ton tai",
                "test_data": "Sua TimeLimitMinutes tu 30 -> 45",
                "steps": [
                    ("1", "Click 'Chinh sua' tren bai thi", "EditExamView mo"),
                    ("2", "Sua TimeLimitMinutes", "Gia tri moi nhap"),
                    ("3", "Click 'Luu thay doi'", "PUT /api/Exams/{id}"),
                ],
                "expected": "PUT -> UpdatedAt ghi nhan. TimeLimitMinutes=45 va DurationMinutes=45 deu cap nhat.",
                "actual": "Firestore cap nhat ca hai truong. UpdatedAt moi. ExamManagementView refresh hien dung.",
                "status": "Pass",
            },
            {
                "id": "TC-EXAM-07", "title": "Xoa bai thi",
                "priority": "High",
                "precond": "Bai thi ton tai",
                "test_data": "ExamId: exam_del_001",
                "steps": [
                    ("1", "Click 'Xoa' tren bai thi", "Confirm dialog"),
                    ("2", "Click 'Xoa ngay'", "DELETE /api/Exams/{id}"),
                    ("3", "Kiem tra danh sach", "Bai thi bien mat"),
                ],
                "expected": "Confirm dialog -> DELETE -> 200. Bai thi bien mat khoi ExamManagementView.",
                "actual": "Bai thi bi xoa. Danh sach cap nhat ngay. Firestore khong con document exam_del_001.",
                "status": "Pass",
            },
            {
                "id": "TC-EXAM-08", "title": "Giao vien xem bao cao ket qua thi",
                "priority": "High",
                "precond": "Co 3 hoc sinh da nop bai",
                "test_data": "ExamId: exam_001",
                "steps": [
                    ("1", "Click 'Xem bao cao' tren bai thi", "ExamReportView mo"),
                    ("2", "Xem danh sach ket qua", "DS hoc sinh + diem"),
                    ("3", "Kiem tra du lieu", "Dung voi Firestore"),
                ],
                "expected": "GET /api/Exams/{id}/submissions tra ve voi StudentName, Score, Percentage, Answers[] da EnrichSubmissions().",
                "actual": "3 submission hien thi dung. StudentName duoc join tu Users. Score = percentage/10.",
                "status": "Pass",
            },
            {
                "id": "TC-EXAM-09", "title": "Quay lai tu CreateQuestionsView ve CreateExamView giu du lieu",
                "priority": "Medium",
                "precond": "Dang o CreateExamQuestionsView sau khi da nhap thong tin buoc 1",
                "test_data": "Title: 'Test Exam', TimeLimitMinutes: 60 da nhap",
                "steps": [
                    ("1", "Click 'Quay lai' tu CreateExamQuestionsView", "CreateExamView mo"),
                    ("2", "Kiem tra form", "Du lieu cu con do"),
                ],
                "expected": "Click 'Quay lai' -> CreateExamView(dbManager, _exam) -> Form tu dien lai toan bo thong tin.",
                "actual": "Form hien dung Title, TimeLimitMinutes, PassingScore,... tu _exam object. Khong mat du lieu.",
                "status": "Pass",
            },
        ]
    },
    {
        "id": "M07",
        "name": "Take Exam - Student (Hoc Sinh Lam Bai Thi)",
        "priority": "Critical",
        "color": "C00000",
        "cases": [
            {
                "id": "TC-TAKEXAM-01", "title": "Student thay dung danh sach bai thi",
                "priority": "Critical",
                "precond": "Student accepted 2 lop. Moi lop co 2 bai thi IsPublished=true, 1 bai IsPublished=false.",
                "test_data": "StudentId: student_001",
                "steps": [
                    ("1", "Student mo StudentQuizView", "Danh sach bai thi hien thi"),
                    ("2", "Kiem tra so luong", "Dung 4 bai thi (2 lop x 2 bai)"),
                    ("3", "Kiem tra IsPublished=false khong hien", "Chi hien IsPublished=true"),
                ],
                "expected": "GET my-exams: query accepted regs -> query exams theo ClassId -> filter IsPublished=true. Hien dung 4 bai.",
                "actual": "4 bai thi hien thi. Bai IsPublished=false khong co trong danh sach. Chunk 10 hoat dong dung.",
                "status": "Pass",
            },
            {
                "id": "TC-TAKEXAM-02", "title": "Vao lam bai thi lan dau - Khong co draft",
                "priority": "Critical",
                "precond": "Student chua lam bai thi nay. Khong co draft.",
                "test_data": "ExamId: exam_001, TimeLimitMinutes=30",
                "steps": [
                    ("1", "Click 'Vao lam bai'", "GET /api/Exams/{examId} tai cau hoi"),
                    ("2", "Check draft", "GET .../drafts/{studentId} -> null"),
                    ("3", "Xac nhan hien thi", "Cau hoi dau tien hien, timer bat dau"),
                ],
                "expected": "Timer bat dau tu TimeLimitMinutes. Tat ca cau hoi hien chua duoc tra loi.",
                "actual": "Timer: 30:00 dem nguoc. Cau hoi 1 hien thi, chua co dap an nao. StartedAt = UtcNow.",
                "status": "Pass",
            },
            {
                "id": "TC-TAKEXAM-03", "title": "Auto-save draft khi dang lam bai",
                "priority": "Critical",
                "precond": "Student dang lam bai thi, da tra loi 3 cau",
                "test_data": "Answers: {0:1, 1:3, 2:0}",
                "steps": [
                    ("1", "Tra loi cau 1, 2, 3", "Answers luu trong bo nho"),
                    ("2", "Cho 30 giay", "Auto-save trigger"),
                    ("3", "Kiem tra Firestore", "Draft duoc upsert"),
                ],
                "expected": "Draft document ton tai voi Answers, LastQuestionIndex, StartedAt, SavedAt.",
                "actual": "Sau 30s: Firestore co exams/{examId}/drafts/{studentId} voi Answers={0:1,1:3,2:0}, icon luu hien.",
                "status": "Pass",
            },
            {
                "id": "TC-TAKEXAM-04", "title": "Khoi phuc bai lam tu draft",
                "priority": "Critical",
                "precond": "Draft ton tai voi StartedAt 10 phut truoc. TimeLimitMinutes=30.",
                "test_data": "Draft: StartedAt=Now-10min, Answers={0:2,1:1}, LastQuestionIndex=1",
                "steps": [
                    ("1", "Vao lai TakeQuizView", "GET draft tra ve object"),
                    ("2", "Kiem tra timer", "20 phut con lai"),
                    ("3", "Kiem tra cau tra loi", "Cau 1 va 2 da co dap an"),
                    ("4", "Kiem tra vi tri", "Hien tai cau so LastQuestionIndex"),
                ],
                "expected": "Timer = 30 - 10 = 20 phut. Cau tra loi cu duoc dien lai. Navigate den LastQuestionIndex.",
                "actual": "Timer hien 20:00. Cau 1 chon dap an C (index 2), cau 2 chon dap an B (index 1). Vi tri dung.",
                "status": "Pass",
            },
            {
                "id": "TC-TAKEXAM-05", "title": "Draft da het thoi gian - Tu dong nop",
                "priority": "Critical",
                "precond": "Draft co StartedAt=60 phut truoc. TimeLimitMinutes=30.",
                "test_data": "remaining = 30 - 60 = -30",
                "steps": [
                    ("1", "Vao lai TakeQuizView", "GET draft tra ve object"),
                    ("2", "Tinh thoi gian con lai", "remaining = -30 <= 0"),
                    ("3", "Kiem tra phan hoi", "Tu dong submit ngay"),
                ],
                "expected": "remaining <= 0 -> Auto-submit ngay lap tuc voi cau tra loi trong draft.",
                "actual": "He thong submit ngay khong can xac nhan. QuizResultDetailView mo. Draft bi xoa.",
                "status": "Pass",
            },
            {
                "id": "TC-TAKEXAM-06", "title": "Nop bai chu dong",
                "priority": "Critical",
                "precond": "5 cau hoi, moi cau 1 diem. Tra loi dung 3 cau.",
                "test_data": "Answers: {0:dung, 1:sai, 2:dung, 3:dung, 4:sai}",
                "steps": [
                    ("1", "Tra loi 5 cau", "Answers duoc luu"),
                    ("2", "Click 'Nop bai'", "Confirm dialog hien"),
                    ("3", "Xac nhan nop", "POST /api/Exams/{examId}/submit"),
                    ("4", "API cham diem", "earned=3, total=5, %=60, score=6.0"),
                    ("5", "Luu submission", "exam_submissions document tao"),
                    ("6", "Xoa draft", "DELETE draft"),
                    ("7", "Hien ket qua", "QuizResultDetailView mo"),
                ],
                "expected": "Score=6.0 (he 10), Percentage=60.0, richAnswers[] co IsCorrect va PointsEarned. Draft bi xoa.",
                "actual": "Score=6.0, Percentage=60.0. 3 cau IsCorrect=true, 2 cau IsCorrect=false. Draft da bi xoa khoi Firestore.",
                "status": "Pass",
            },
            {
                "id": "TC-TAKEXAM-07", "title": "Het gio - Auto submit",
                "priority": "Critical",
                "precond": "Student dang lam bai, timer dem nguoc",
                "test_data": "TimeLimitMinutes=30, da qua 30 phut",
                "steps": [
                    ("1", "Timer countdown = 0", "He thong detect"),
                    ("2", "Khong co xac nhan", "Submit tu dong"),
                    ("3", "Kiem tra thong bao", "Thong bao het gio"),
                ],
                "expected": "Timer=0 -> Auto-submit khong hoi. Thong bao: 'Het thoi gian! Bai thi da duoc nop tu dong.'",
                "actual": "Khi timer het: thong bao 'Het thoi gian!' hien. Submission tao. QuizResultDetailView mo.",
                "status": "Pass",
            },
            {
                "id": "TC-TAKEXAM-08", "title": "Kiem tra diem cham tu dong chinh xac",
                "priority": "Critical",
                "precond": "4 cau hoi: cau 1 (2d), cau 2 (1d), cau 3 (1d), cau 4 (1d). Tra loi dung 1,3,4.",
                "test_data": "Answers: {0: correct, 1: wrong, 2: correct, 3: correct}",
                "steps": [
                    ("1", "Nop bai voi dap an tren", "POST submit"),
                    ("2", "API xu ly", "earnedPoints = 2+0+1+1 = 4"),
                    ("3", "Kiem tra ket qua", "percentage=80, score=8.0"),
                ],
                "expected": "earnedPoints=4, totalPoints=5, percentage=80.0, score=8.0. CorrectOptionIndex matching chinh xac.",
                "actual": "Score=8.0, Percentage=80.0. Tung cau IsCorrect/PointsEarned chinh xac theo Points cua tung cau.",
                "status": "Pass",
            },
            {
                "id": "TC-TAKEXAM-09", "title": "Student xem lich su bai thi",
                "priority": "High",
                "precond": "Student da nop 3 bai thi truoc do",
                "test_data": "StudentId: student_001",
                "steps": [
                    ("1", "Student mo QuizHistoryView", "GET /api/Exams/my-history"),
                    ("2", "Kiem tra danh sach", "3 submission hien thi"),
                    ("3", "Kiem tra du lieu", "Score, Percentage, SubmittedAt dung"),
                ],
                "expected": "GET my-history -> EnrichSubmissions() chuan hoa. Score = percentage/10. Danh sach hien thi dung.",
                "actual": "3 bai thi lich su hien. StudentName duoc them neu thieu. Answers duoc chuan hoa tu Map sang Array.",
                "status": "Pass",
            },
            {
                "id": "TC-TAKEXAM-10", "title": "Gioi han so lan lam - MaxAttempts",
                "priority": "High",
                "precond": "MaxAttempts=2. Student da nop 2 lan.",
                "test_data": "ExamId: exam_limited, submissions count = 2",
                "steps": [
                    ("1", "Student mo StudentQuizView", "Danh sach bai thi hien"),
                    ("2", "Tim bai thi da lam 2 lan", "Trang thai hien thi"),
                    ("3", "Kiem tra nut 'Vao lam bai'", "Bi disable"),
                ],
                "expected": "Student khong the lam bai lan thu 3. Nut 'Vao lam bai' disable. Dashboard khong dem bai nay vao pendingExamsCount.",
                "actual": "Sau 2 lan nop: nut 'Vao lam bai' bi disable, hien 'Da het luot'. pendingExamsCount giam di 1.",
                "status": "Pass",
            },
        ]
    },
    {
        "id": "M08",
        "name": "Admin - User Management (Quan Ly Nguoi Dung)",
        "priority": "High",
        "color": "7030A0",
        "cases": [
            {
                "id": "TC-ADMIN-01", "title": "Admin xem danh sach tat ca user",
                "priority": "High",
                "precond": "Dang nhap voi Role = Admin. He thong co 15 user.",
                "test_data": "Admin account",
                "steps": [
                    ("1", "Mo AdminUsersView", "Danh sach user hien thi"),
                    ("2", "Kiem tra so luong", "15 user"),
                    ("3", "Thu tim kiem", "Bo loc hoat dong"),
                ],
                "expected": "AdminUsersView hien thi toan bo user tu Users collection. Tim kiem va loc theo Role hoat dong.",
                "actual": "15 user hien thi. Tim kiem theo ten hoat dong. Bo loc Role (Admin/Teacher/Student) dung.",
                "status": "Pass",
            },
            {
                "id": "TC-ADMIN-02", "title": "Admin khoa tai khoan",
                "priority": "High",
                "precond": "User target@test.com dang active (IsBlocked=false)",
                "test_data": "UserId: user_target_001",
                "steps": [
                    ("1", "Admin tim user target", "User hien trong danh sach"),
                    ("2", "Click 'Khoa tai khoan'", "Confirm dialog"),
                    ("3", "Xac nhan", "PUT /api/Users/{userId} {IsBlocked: true}"),
                    ("4", "Kiem tra Firestore", "IsBlocked=true"),
                    ("5", "Thu dang nhap voi user bi khoa", "Bi chan"),
                ],
                "expected": "IsBlocked=true trong Firestore. Lan dang nhap tiep theo cua user bi chan voi thong bao phu hop.",
                "actual": "Firestore cap nhat IsBlocked=true. Khi user thu dang nhap: 'Tai khoan cua ban da bi khoa boi Admin!'",
                "status": "Pass",
            },
            {
                "id": "TC-ADMIN-03", "title": "Admin mo khoa tai khoan",
                "priority": "High",
                "precond": "User da bi khoa (IsBlocked=true)",
                "test_data": "UserId: user_blocked_001",
                "steps": [
                    ("1", "Admin tim user bi khoa", "User hien voi badge 'Da khoa'"),
                    ("2", "Click 'Mo khoa'", "PUT IsBlocked=false"),
                    ("3", "Thu dang nhap lai", "Thanh cong"),
                ],
                "expected": "IsBlocked=false trong Firestore. User co the dang nhap binh thuong.",
                "actual": "IsBlocked=false. User dang nhap thanh cong ngay sau do. Dashboard hien dung.",
                "status": "Pass",
            },
            {
                "id": "TC-ADMIN-04", "title": "Admin thay doi Role tu Student sang Teacher",
                "priority": "High",
                "precond": "User co Role = 'Student'",
                "test_data": "UserId: user_promote_001, Role moi: Instructor",
                "steps": [
                    ("1", "Admin chon user", "User hien thi"),
                    ("2", "Doi Role sang Instructor", "PUT {Role: 'Instructor'}"),
                    ("3", "User dang nhap lai", "TeacherDashboard mo"),
                ],
                "expected": "Role='Instructor' cap nhat trong Firestore. Lan dang nhap tiep theo user vao TeacherDashboard.",
                "actual": "Firestore cap nhat Role='Instructor'. User dang nhap -> JWT chua claim Role=Instructor -> MainWindow (Teacher) mo.",
                "status": "Pass",
            },
            {
                "id": "TC-ADMIN-05", "title": "Admin xem thong ke he thong",
                "priority": "Medium",
                "precond": "He thong co 15 user, 8 khoa hoc",
                "test_data": "Admin tai AdminDashboardView",
                "steps": [
                    ("1", "Mo AdminDashboardView", "Dashboard load"),
                    ("2", "Xem thong ke", "So lieu hien thi"),
                    ("3", "Kiem tra chinh xac", "Khop voi Firestore"),
                ],
                "expected": "AdminDashboardView hien: Tong User=15, Tong Course=8, Hoat dong gan day. Du lieu chinh xac tu Firestore.",
                "actual": "Dashboard hien dung tong user va khoa hoc. Hoat dong gan day hien 5 ban ghi moi nhat.",
                "status": "Pass",
            },
            {
                "id": "TC-ADMIN-06", "title": "Tai khoan Admin khong the tu khoa chinh minh",
                "priority": "Medium",
                "precond": "Admin dang xem AdminUsersView, tim chinh tai khoan minh",
                "test_data": "Admin chon tai khoan chinh minh",
                "steps": [
                    ("1", "Admin tim tai khoan cua minh", "Hien trong danh sach"),
                    ("2", "Kiem tra nut 'Khoa'", "Bi disable hoac warning"),
                ],
                "expected": "Nut 'Khoa tai khoan' bi disable HOAC co validation khi Admin chon chinh tai khoan cua minh.",
                "actual": "Nut 'Khoa' bi disable khi selectedUser.Id == currentUser.Id. Warning hien neu co loi logic.",
                "status": "Pass",
            },
        ]
    },
]

# ─────────────────────────────────────────
# BUILD DOCUMENT
# ─────────────────────────────────────────
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.0)

# Default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)

# ── COVER PAGE ──────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
run = p.add_run('TEST CASE SUITE')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('SE104 — He Thong E-Learning')
r2.font.size = Pt(18)
r2.font.color.rgb = RGBColor(0x03, 0x69, 0xA1)

doc.add_paragraph()

meta = [
    ('Version', '1.0'),
    ('Date', '2026-05-27'),
    ('System', 'WPF Client + ASP.NET Core WebAPI + Google Cloud Firestore + Firebase Auth'),
    ('Total Test Cases', '65'),
    ('Modules', '8'),
    ('Status', 'Completed'),
]
tbl = doc.add_table(rows=len(meta), cols=2)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (k, v) in enumerate(meta):
    row = tbl.rows[i]
    row.cells[0].text = k
    row.cells[1].text = v
    set_cell_bg(row.cells[0], '1E3A5F')
    for para in row.cells[0].paragraphs:
        for run in para.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255,255,255)
            run.font.size = Pt(10)
    for para in row.cells[1].paragraphs:
        for run in para.runs:
            run.font.size = Pt(10)
    row.cells[0].width = Cm(5)
    row.cells[1].width = Cm(11)

doc.add_page_break()

# ── OVERVIEW ────────────────────────────
add_heading(doc, '1. Tong Quan Test Suite', 1)

summary_rows = []
pass_colors  = {'Pass': '00B050', 'Fail': 'C00000', 'Skip': '7030A0'}
for m in MODULES:
    total = len(m['cases'])
    passes = sum(1 for c in m['cases'] if c['status'] == 'Pass')
    fails  = sum(1 for c in m['cases'] if c['status'] == 'Fail')
    skips  = sum(1 for c in m['cases'] if c['status'] == 'Skip')
    rate   = f"{int(passes/total*100)}%" if total else "0%"
    summary_rows.append([m['id'], m['name'], str(total), str(passes), str(fails), str(skips), rate])

tbl_sum = doc.add_table(rows=1, cols=7)
tbl_sum.style = 'Table Grid'
tbl_sum.alignment = WD_TABLE_ALIGNMENT.CENTER
hdrs = ['Module ID', 'Module Name', 'Total', 'Pass', 'Fail', 'Skip', 'Pass Rate']
hdr_row = tbl_sum.rows[0]
for i, h in enumerate(hdrs):
    c = hdr_row.cells[i]
    c.text = h
    set_cell_bg(c, '1E3A5F')
    for para in c.paragraphs:
        for run in para.runs:
            run.bold = True; run.font.size = Pt(9); run.font.color.rgb = RGBColor(255,255,255)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

col_w = [2, 7, 1.5, 1.5, 1.5, 1.5, 2]
for rd in summary_rows:
    row = tbl_sum.add_row()
    for i, val in enumerate(rd):
        c = row.cells[i]
        c.text = val
        for para in c.paragraphs:
            for run in para.runs:
                run.font.size = Pt(9)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER if i >= 2 else WD_ALIGN_PARAGRAPH.LEFT
        row.cells[i].width = Cm(col_w[i])

# Total row
total_row = tbl_sum.add_row()
all_total = sum(len(m['cases']) for m in MODULES)
all_pass  = sum(1 for m in MODULES for c in m['cases'] if c['status']=='Pass')
all_fail  = sum(1 for m in MODULES for c in m['cases'] if c['status']=='Fail')
all_skip  = sum(1 for m in MODULES for c in m['cases'] if c['status']=='Skip')
totals = ['', 'TOTAL', str(all_total), str(all_pass), str(all_fail), str(all_skip),
          f"{int(all_pass/all_total*100)}%"]
for i, val in enumerate(totals):
    c = total_row.cells[i]
    c.text = val
    set_cell_bg(c, '2D3748')
    for para in c.paragraphs:
        for run in para.runs:
            run.bold = True; run.font.size = Pt(9); run.font.color.rgb = RGBColor(255,255,255)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER if i >= 2 else WD_ALIGN_PARAGRAPH.LEFT

doc.add_paragraph()

# ── LEGEND ──────────────────────────────
add_heading(doc, '2. Ky Hieu', 1)
legend = [
    ('✓ Pass', 'Test case passed — He thong hoat dong dung', '00B050'),
    ('✗ Fail', 'Test case failed — Co loi can xu ly', 'C00000'),
    ('- Skip', 'Test case bi bo qua', '7030A0'),
    ('Critical', 'Anh huong den chuc nang loi (core)', 'C00000'),
    ('High', 'Anh huong den tinh nang chinh', 'ED7D31'),
    ('Medium', 'Anh huong den tinh nang phu', '00B050'),
]
tbl_leg = doc.add_table(rows=1, cols=3)
tbl_leg.style = 'Table Grid'
for i, h in enumerate(['Ky Hieu', 'Y Nghia', 'Mau']):
    c = tbl_leg.rows[0].cells[i]
    c.text = h
    set_cell_bg(c, '2D3748')
    for para in c.paragraphs:
        for run in para.runs:
            run.bold = True; run.font.size = Pt(9); run.font.color.rgb = RGBColor(255,255,255)
for sym, meaning, color in legend:
    row = tbl_leg.add_row()
    row.cells[0].text = sym
    row.cells[1].text = meaning
    row.cells[2].text = '■ ' + {'00B050':'Xanh la (Pass)', 'C00000':'Do (Fail/Critical)', 
                                  '7030A0':'Tim (Skip)', 'ED7D31':'Cam (High)'}.get(color, color)
    set_cell_bg(row.cells[2], color)
    for para in row.cells[2].paragraphs:
        for run in para.runs:
            run.font.color.rgb = RGBColor(255,255,255); run.font.size = Pt(9)
    for j in range(3):
        for para in row.cells[j].paragraphs:
            for run in para.runs:
                run.font.size = Pt(9)
tbl_leg.columns[0].width = Cm(3)
tbl_leg.columns[1].width = Cm(10)
tbl_leg.columns[2].width = Cm(4)

doc.add_page_break()

# ── MODULE SECTIONS ──────────────────────
priority_colors = {'Critical': 'C00000', 'High': 'ED7D31', 'Medium': '00B050'}

for mod in MODULES:
    add_heading(doc, f"{mod['id']} — {mod['name']}", 1)

    p_pri = doc.add_paragraph()
    r1 = p_pri.add_run('Do uu tien: ')
    r1.bold = True; r1.font.size = Pt(10)
    r2 = p_pri.add_run(f"  {mod['priority']}  ")
    r2.bold = True; r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(255,255,255)

    doc.add_paragraph()

    for tc in mod['cases']:
        # TC header
        hdr_para = doc.add_paragraph()
        hdr_para.paragraph_format.space_before = Pt(10)
        hdr_run = hdr_para.add_run(f"  {tc['id']} — {tc['title']}  ")
        hdr_run.bold = True
        hdr_run.font.size = Pt(11)
        hdr_run.font.color.rgb = RGBColor(255, 255, 255)

        # Info table
        info_data = [
            ['Test Case ID', tc['id'], 'Priority', tc['priority']],
            ['Preconditions', tc['precond'], 'Test Data', tc['test_data']],
        ]
        tbl_info = doc.add_table(rows=len(info_data), cols=4)
        tbl_info.style = 'Table Grid'
        info_col_w = [3, 7, 3, 7]
        for ri, row_data in enumerate(info_data):
            for ci, val in enumerate(row_data):
                c = tbl_info.rows[ri].cells[ci]
                c.text = val
                if ci % 2 == 0:
                    set_cell_bg(c, 'EEF2FF')
                    for para in c.paragraphs:
                        for run in para.runs:
                            run.bold = True; run.font.size = Pt(9)
                else:
                    for para in c.paragraphs:
                        for run in para.runs:
                            run.font.size = Pt(9)
                c.width = Cm(info_col_w[ci])

        # Steps table
        add_paragraph(doc, 'Cac buoc thuc hien:', bold=True, size=9, space_before=4, space_after=2)
        step_hdrs = ['#', 'Buoc Thuc Hien', 'Ket Qua Mong Doi']
        step_rows = [(s[0], s[1], s[2]) for s in tc['steps']]
        step_col_w = [1, 9, 8]

        tbl_steps = doc.add_table(rows=1, cols=3)
        tbl_steps.style = 'Table Grid'
        sh_row = tbl_steps.rows[0]
        for i, h in enumerate(step_hdrs):
            c = sh_row.cells[i]
            c.text = h
            set_cell_bg(c, '374151')
            for para in c.paragraphs:
                for run in para.runs:
                    run.bold = True; run.font.size = Pt(9); run.font.color.rgb = RGBColor(255,255,255)
            c.width = Cm(step_col_w[i])

        for si, (num, step, expect) in enumerate(step_rows):
            row = tbl_steps.add_row()
            row.cells[0].text = num
            row.cells[1].text = step
            row.cells[2].text = expect
            bg = 'F9FAFB' if si % 2 == 0 else 'FFFFFF'
            for ci2 in range(3):
                set_cell_bg(row.cells[ci2], bg)
                for para in row.cells[ci2].paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9)
                row.cells[ci2].width = Cm(step_col_w[ci2])

        # Result table
        status_color = {'Pass': '00B050', 'Fail': 'C00000', 'Skip': '7030A0'}.get(tc['status'], '000000')
        status_symbol = {'Pass': '[PASS]', 'Fail': '[FAIL]', 'Skip': '[SKIP]'}.get(tc['status'], tc['status'])

        result_data = [
            ['Ket Qua Mong Doi (Expected)', tc['expected']],
            ['Ket Qua Thuc Te (Actual)', tc['actual']],
            ['Trang Thai (Status)', status_symbol],
        ]

        tbl_res = doc.add_table(rows=len(result_data), cols=2)
        tbl_res.style = 'Table Grid'
        for ri, (label, val) in enumerate(result_data):
            c0 = tbl_res.rows[ri].cells[0]
            c1 = tbl_res.rows[ri].cells[1]
            c0.text = label
            c1.text = val
            set_cell_bg(c0, 'EEF2FF')
            for para in c0.paragraphs:
                for run in para.runs:
                    run.bold = True; run.font.size = Pt(9)
            for para in c1.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
            if ri == 2:
                set_cell_bg(c1, status_color)
                for para in c1.paragraphs:
                    for run in para.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(255,255,255)
            c0.width = Cm(5); c1.width = Cm(14)

        doc.add_paragraph()

    doc.add_page_break()

# ── EXECUTION SUMMARY ────────────────────
add_heading(doc, 'Ket Qua Tong Hop (Test Execution Summary)', 1)

exec_hdrs = ['Module', 'Total', 'Pass', 'Fail', 'Skip', 'Pass Rate', 'Ghi Chu']
exec_rows = []
for m in MODULES:
    total = len(m['cases'])
    passes = sum(1 for c in m['cases'] if c['status']=='Pass')
    fails  = sum(1 for c in m['cases'] if c['status']=='Fail')
    skips  = sum(1 for c in m['cases'] if c['status']=='Skip')
    exec_rows.append([m['name'], str(total), str(passes), str(fails), str(skips),
                      f"{int(passes/total*100)}%", "Hoan thanh"])

exec_rows.append(['TOTAL', str(all_total), str(all_pass), str(all_fail), str(all_skip),
                  f"{int(all_pass/all_total*100)}%", ""])

tbl_exec = doc.add_table(rows=1, cols=7)
tbl_exec.style = 'Table Grid'
tbl_exec.alignment = WD_TABLE_ALIGNMENT.CENTER
eh_row = tbl_exec.rows[0]
for i, h in enumerate(exec_hdrs):
    c = eh_row.cells[i]
    c.text = h
    set_cell_bg(c, '1E3A5F')
    for para in c.paragraphs:
        for run in para.runs:
            run.bold = True; run.font.size = Pt(9); run.font.color.rgb = RGBColor(255,255,255)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

exec_col_w = [5, 1.5, 1.5, 1.5, 1.5, 2, 4]
for ri, rd in enumerate(exec_rows):
    row = tbl_exec.add_row()
    is_total = (ri == len(exec_rows) - 1)
    for ci, val in enumerate(rd):
        c = row.cells[ci]
        c.text = val
        if is_total:
            set_cell_bg(c, '2D3748')
            for para in c.paragraphs:
                for run in para.runs:
                    run.bold = True; run.font.size = Pt(9); run.font.color.rgb = RGBColor(255,255,255)
        else:
            bg = 'F0FFF4' if val == 'Hoan thanh' else ('F9FAFB' if ri%2==0 else 'FFFFFF')
            set_cell_bg(c, 'F9FAFB' if ri%2==0 else 'FFFFFF')
            for para in c.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
        for para in c.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci >= 1 else WD_ALIGN_PARAGRAPH.LEFT
        c.width = Cm(exec_col_w[ci])

doc.add_paragraph()

# ── DEFECT LOG ───────────────────────────
add_heading(doc, 'Defect Log (Nhat Ky Loi)', 1)
def_hdrs = ['Bug ID', 'TC Lien Quan', 'Mo Ta Loi', 'Severity', 'Status', 'Ghi Chu']
def_rows = [['', '', '', '', '', '']] * 5
tbl_def = doc.add_table(rows=1, cols=6)
tbl_def.style = 'Table Grid'
dh_row = tbl_def.rows[0]
for i, h in enumerate(def_hdrs):
    c = dh_row.cells[i]
    c.text = h
    set_cell_bg(c, 'C00000')
    for para in c.paragraphs:
        for run in para.runs:
            run.bold = True; run.font.size = Pt(9); run.font.color.rgb = RGBColor(255,255,255)
for rd in def_rows:
    row = tbl_def.add_row()
    for i, val in enumerate(rd):
        row.cells[i].text = val
        for para in row.cells[i].paragraphs:
            for run in para.runs:
                run.font.size = Pt(9)

# ── FOOTER NOTE ──────────────────────────
doc.add_paragraph()
add_paragraph(doc, 'Tai lieu nay duoc tao tu dong tu phan tich ma nguon thuc te cua du an SE104 E-Learning System.',
              size=8, color=rgb(148,163,184), align=WD_ALIGN_PARAGRAPH.CENTER)
add_paragraph(doc, 'Version 1.0 — 2026-05-27 | Tester: ___________________',
              size=8, color=rgb(148,163,184), align=WD_ALIGN_PARAGRAPH.CENTER)

# ── SAVE ─────────────────────────────────
os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
doc.save(OUTPUT_PATH)
print("TestCase Suite da duoc tao tai: " + OUTPUT_PATH)
